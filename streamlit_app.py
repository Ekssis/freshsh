"""
FRESH - Генератор документов с искусственным ПВ
v6 FINAL — исправлены: даты/суммы, сохранение таблицы ном.док,
           перенос ФИО, парсинг марки авто
"""

import streamlit as st
import re, sys, subprocess
from datetime import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    import num2words as nw
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "num2words", "-q"])
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    import num2words as nw

# ─────────────────────────────────────────────────────────────
# УТИЛИТЫ СУММ
# ─────────────────────────────────────────────────────────────

def _rub_word(n: int) -> str:
    n100, n10 = n % 100, n % 10
    if 11 <= n100 <= 19: return "рублей"
    if n10 == 1:         return "рубль"
    if 2 <= n10 <= 4:    return "рубля"
    return "рублей"

def format_amount(amount: float) -> str:
    rub = int(amount)
    kop = round((amount - rub) * 100)
    return f"{rub:,}".replace(",", " ") + f",{kop:02d}"

def amount_to_words(amount: float) -> str:
    rub = int(amount)
    kop = round((amount - rub) * 100)
    w = nw.num2words(rub, lang="ru")
    return w[0].upper() + w[1:] + f" {_rub_word(rub)} {kop:02d} копеек"

def parse_amount(text: str) -> float:
    if not text: return 0.0
    cleaned = re.sub(r"[^\d,.]", "", text.strip()).replace(",", ".")
    parts = cleaned.split(".")
    if len(parts) > 2:
        cleaned = "".join(parts[:-1]) + "." + parts[-1]
    try:    return float(cleaned)
    except: return 0.0

# ─────────────────────────────────────────────────────────────
# DOCX УТИЛИТЫ
# ─────────────────────────────────────────────────────────────

RE_DATE = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")

# ── ИСПРАВЛЕНИЕ #1: RE_AMOUNT больше не матчит даты ──
# Добавлен negative lookahead (?!\.\d{2}) — блокирует совпадение
# с датами формата DD.MM.YYYY (где после DD.MM идёт .YYYY)
RE_AMOUNT = re.compile(
    r"(?<!\d)"                              # не после цифры
    r"(\d{1,3}(?:[ \u00A0]\d{3})*"          # целая часть с разделителями тысяч
    r"[.,]\d{2})"                           # десятичная часть (,XX или .XX)
    r"(?!\.\d{2})"                          # НЕ дата: после суммы нет .XX
    r"(?!\d)"                               # не перед цифрой
)

RE_WORDS = re.compile(
    r"[А-ЯЁа-яё][а-яёА-ЯЁ\s\-\,]+"
    r"(?:тысяч|миллион|миллиард|рубл)[а-яё\s\-\,]*"
    r"\d{2}\s+копеек\s*"
)

def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)

def _set_para(para, text: str):
    """Пишет text в параграф, очищая все runs кроме первого."""
    if para is None: return
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(text)

def _remove_nds(text: str) -> str:
    text = re.sub(r",?\s*в\s+т\.?\s*ч\.?\s*НДС[^.;\n]*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\(?22/122%?\)?\s*[\d\s,.]+руб\.?",   "", text, flags=re.IGNORECASE)
    return re.sub(r"\s+", " ", text).strip()

def _iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            if hf:
                yield from hf.paragraphs

def insert_paragraph_after(paragraph, text, source_para=None):
    new_p = OxmlElement("w:p")
    paragraph._p.getparent().insert(
        paragraph._p.getparent().index(paragraph._p) + 1, new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.font.size = Pt(8)
    if source_para and source_para.runs:
        run.font.name = source_para.runs[0].font.name
    new_para.paragraph_format.space_before = Pt(0)
    new_para.paragraph_format.space_after  = Pt(4)
    new_para.paragraph_format.line_spacing = 1.0
    return new_para

# ── ИСПРАВЛЕНИЕ #3: Расширение ячейки и отключение переноса ──
def expand_cell_no_wrap(cell, width_twips=9000):
    """Расширяет ячейку таблицы и отключает перенос слов."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Удаляем старую ширину
    for elem in tcPr.findall(qn('w:tcW')):
        tcPr.remove(elem)
    # Ставим новую
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

    # Удаляем старый noWrap
    for elem in tcPr.findall(qn('w:noWrap')):
        tcPr.remove(elem)
    # Добавляем noWrap — текст не переносится
    noWrap = OxmlElement('w:noWrap')
    tcPr.append(noWrap)

    # Также убираем фиксированную ширину у параграфов
    for para in cell.paragraphs:
        pPr = para._p.get_or_add_pPr()
        # Убираем переносы
        for elem in pPr.findall(qn('w:kinsoku')):
            pPr.remove(elem)

def set_table_autofit(tbl):
    """Переключает таблицу в autofit-режим."""
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)
    for elem in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(elem)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'autofit')
    tblPr.append(layout)

# ─────────────────────────────────────────────────────────────
# ПАРСИНГ СЧЁТА №1
# ─────────────────────────────────────────────────────────────

def extract_invoice_data(doc: Document) -> dict:
    data = {}
    full_text = "\n".join([p.text for p in _iter_all_paragraphs(doc)])

    # 1. Поиск ДКП и даты
    m = re.search(
        r"[Пп]родажа\s+[тТ][/\\][сС]\s+№\s*([\w\-/]+).*?от\s+(\d{2}\.\d{2}\.(?:\d{2}|\d{4}))",
        full_text
    )
    if m:
        data["dkp_number"] = m.group(1).strip()
        data["date"] = m.group(2)

    # 2. Поиск покупателя
    m = re.search(
        r"[Пп]окупатель[:\s]+(?:ИНН\s+\d+[,\s]+)?([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)",
        full_text
    )
    if m: data["buyer_name"] = m.group(1).strip()

    # ── ИСПРАВЛЕНИЕ #4: Улучшенный парсинг марки и модели авто ──
    colors = [
        "серый", "белый", "черный", "чёрный", "синий", "красный",
        "серебристый", "золотистый", "коричневый", "зелёный", "зеленый",
        "бежевый", "жёлтый", "желтый"
    ]
    color_pattern = "|".join(colors)

    for tbl in doc.tables:
        for row in tbl.rows:
            cells_text = [c.text.strip() for c in row.cells]
            rt = " ".join(c for c in cells_text if c)
            if not rt:
                continue

            # Сначала ищем VIN (17 символов, без I, O, Q)
            vin_match = re.search(r"\b([A-HJ-NPR-Z0-9]{17})\b", rt, re.IGNORECASE)
            if not vin_match:
                continue
            vin = vin_match.group(1).upper()

            # Ищем цвет
            color_match = re.search(f"({color_pattern})", rt, re.IGNORECASE)
            if not color_match:
                continue
            color = color_match.group(1).strip()

            # Ищем гос. номер: буква-3цифры-2буквы-регион(2-3 цифры)
            reg_match = re.search(
                r"([АВЕКМНОРСТУХABEKMHOPCTYX]\d{3}[АВЕКМНОРСТУХABEKMHOPCTYX]{2}\s*\d{2,3})",
                rt, re.IGNORECASE
            )
            if not reg_match:
                # Запасной вариант: № <буквенно-цифровой>
                reg_match = re.search(r"№\s*([A-ZА-ЯЁ0-9]+)", rt, re.IGNORECASE)
            reg = reg_match.group(1).strip() if reg_match else ""

            # Марка = весь текст ДО цвета, очищенный от мусора
            brand = rt[:color_match.start()].strip()
            # Убираем ведущие номера строк, пунктуацию
            brand = re.sub(r"^[\d№\s\.\,\;\-]+", "", brand).strip()
            # Убираем возможные заголовки колонок
            brand = re.sub(
                r"^(Товар|Наименование|Наименование\s+товара|Услуга)\s*",
                "", brand, flags=re.IGNORECASE
            ).strip()
            brand = brand.rstrip(",.").strip()

            if brand:
                data["car_brand"] = brand
                data["car_color"] = color
                data["car_reg"]   = reg
                data["car_vin"]   = vin
                break
    return data

# ─────────────────────────────────────────────────────────────
# ПРОЦЕССОР ДКП
# ─────────────────────────────────────────────────────────────

def process_dkp(doc: Document, p: dict) -> Document:
    new_str   = format_amount(p["new_price"])
    new_words = amount_to_words(p["new_price"])
    pv_str    = format_amount(p["pv_amount"])
    pv_words  = amount_to_words(p["pv_amount"])

    pv_found  = False
    target_para = None

    for para in doc.paragraphs:
        full = _para_text(para)
        t    = re.sub(r"\s+", " ", full).strip()
        if not t:
            continue

        if t.startswith("Цена ТС оплачивается Покупателем"):
            target_para = para

        if any(w in t.lower() for w in ["гаранти", "техническая защита", "лимит", "пробег"]):
            continue

        if "первоначальный" in t.lower() or "взнос" in t.lower():
            t2 = RE_AMOUNT.sub(pv_str, t)
            t2 = RE_WORDS.sub(pv_words, t2)
            para.text = ""
            run = para.add_run(t2); run.font.size = Pt(8)
            pv_found = True

        elif any(m in t.lower() for m in ["цена договора", "стоимость тс", "цена тс",
                                            "цену за тс", "стоимость автомобиля", "уплачивает покупатель"]):
            t2 = _remove_nds(t)
            if RE_AMOUNT.search(t2): t2 = RE_AMOUNT.sub(new_str, t2)
            if RE_WORDS.search(t2):  t2 = RE_WORDS.sub(new_words, t2)
            t2 = re.sub(r"\s+", " ", t2).strip()
            if not t2.endswith("."): t2 += "."
            para.text = ""
            run = para.add_run(t2); run.font.size = Pt(8)

    if not pv_found and target_para is not None:
        pv_text = (f"Первоначальный взнос по оплате цены Договора составляет "
                   f"{pv_str} руб ({pv_words}).")
        insert_paragraph_after(target_para, pv_text, source_para=target_para)
    return doc

# ─────────────────────────────────────────────────────────────
# ПРОЦЕССОР ПКО  (v6 — даты и суммы разделены, таблица ном. сохранена)
# ─────────────────────────────────────────────────────────────

def process_pko(doc: Document, p: dict) -> Document:
    pv_str   = format_amount(p["pv_amount"])
    pv_words = amount_to_words(p["pv_amount"])
    date     = p["date"]
    buyer    = p["buyer_name"]

    osnov_full = (f"По ДКП №{p['dkp_number']} от {date} "
                  f"за а/м {p['car_brand']} {p['car_color']} "
                  f"№ {p['car_reg']} VIN {p['car_vin']}")

    # ── ШАГ 1: Обработка таблиц ──────────────────────────
    for tbl in doc.tables:
        # Находим колонку "Сумма" по заголовку
        sum_col = -1
        if tbl.rows:
            for ci, hc in enumerate(tbl.rows[0].cells):
                ht = hc.text.strip().lower().rstrip(",")
                if ht == "сумма":
                    sum_col = ci
                    break

        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip()
                if not ct:
                    continue
                ctl = ct.lower()

                # ── ИСПРАВЛЕНИЕ #2: Сохраняем блок "Номер документа" / "Дата составления" ──
                if "номер документа" in ctl or "дата составления" in ctl:
                    # Только обновляем дату внутри, остальное не трогаем
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if RE_DATE.search(pt):
                            _set_para(para, RE_DATE.sub(date, pt))
                    continue

                # ── ИСПРАВЛЕНИЕ #3: "Принято от" — расширяем ячейку, отключаем перенос ──
                if "принято от" in ctl:
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if "принято от" in pt.lower():
                            new_text = re.sub(
                                r"(Принято от[:\s]*).*",
                                rf"\g<1>{buyer}",
                                pt,
                                flags=re.IGNORECASE
                            )
                            _set_para(para, new_text)
                    # Расширяем ячейку и отключаем перенос слов
                    expand_cell_no_wrap(cell, width_twips=9000)
                    set_table_autofit(tbl)
                    continue

                # ── "Основание" / "По ДКП" ──
                if "по дкп" in ctl or "за а/м" in ctl or "основание" in ctl:
                    first_set = False
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if pt.strip() and not first_set:
                            _set_para(para, osnov_full)
                            first_set = True
                        elif first_set:
                            _set_para(para, "")
                    continue

                # ── "Сумма прописью" ──
                if "сумма прописью" in ctl or "прописью" in ctl:
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if RE_WORDS.search(pt):
                            _set_para(para, pv_words)
                        elif RE_AMOUNT.search(pt):
                            _set_para(para, pv_str)
                    continue

                # ── Колонка "Сумма" (по заголовку) — только цифры ──
                if ci == sum_col and ri > 0:
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if RE_AMOUNT.search(pt):
                            _set_para(para, RE_AMOUNT.sub(pv_str, pt))
                    continue

                # ── Ячейка с "Сумма" и числом (не заголовок колонки) ──
                if "сумма" in ctl and RE_AMOUNT.search(ct):
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if RE_AMOUNT.search(pt):
                            new_text = RE_AMOUNT.sub(pv_str, pt)
                            if RE_WORDS.search(new_text):
                                new_text = RE_WORDS.sub(pv_words, new_text)
                            _set_para(para, new_text)
                    continue

                # ── Дата в отдельной ячейке (только дата) ──
                if re.match(r"^\d{2}\.\d{2}\.\d{4}$", ct):
                    if cell.paragraphs:
                        _set_para(cell.paragraphs[0], date)
                    continue

                # ── Очистка НДС (точечно, НЕ удаляя ячейку целиком) ──
                if ("ндс" in ctl or "22/122" in ctl):
                    # Не трогаем ячейки с "номер" или "дата"
                    if "номер" in ctl or "дата" in ctl:
                        continue
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        cleaned = _remove_nds(pt)
                        if cleaned != pt:
                            _set_para(para, cleaned)

    # ── ШАГ 2: Обработка параграфов вне таблиц ──────────
    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        tl = t.lower()

        # ── ИСПРАВЛЕНИЕ #2: Сохраняем "Номер документа" / "Дата составления" ──
        if "номер документа" in tl or "дата составления" in tl:
            if RE_DATE.search(t):
                _set_para(para, RE_DATE.sub(date, t))
            continue

        # Дата в чистом виде
        if re.match(r"^\d{2}\.\d{2}\.\d{4}$", t):
            _set_para(para, date)
            continue

        # "Принято от"
        if "принято от" in tl:
            new_text = re.sub(
                r"(Принято от[:\s]*).*",
                rf"\g<1>{buyer}",
                t,
                flags=re.IGNORECASE
            )
            _set_para(para, new_text)
            continue

        # "По ДКП" / "Основание"
        if "по дкп" in tl or "за а/м" in tl:
            _set_para(para, osnov_full)
            continue

        # "Сумма прописью"
        if "сумма прописью" in tl or "прописью" in tl:
            if RE_WORDS.search(t):
                _set_para(para, pv_words)
            elif RE_AMOUNT.search(t):
                _set_para(para, pv_str)
            continue

        # "Сумма" с числом — заменяем ТОЛЬКО сумму, даты не трогаем
        if "сумма" in tl and RE_AMOUNT.search(t):
            new_text = RE_AMOUNT.sub(pv_str, t)
            if RE_WORDS.search(new_text):
                new_text = RE_WORDS.sub(pv_words, new_text)
            _set_para(para, new_text)
            continue

        # НДС — точечная очистка, не удаление
        if "ндс" in tl or "22/122" in tl:
            if "номер" in tl or "дата" in tl:
                continue
            cleaned = _remove_nds(t)
            if cleaned != t:
                _set_para(para, cleaned)
            continue

    # ── ШАГ 3: Колонтитулы — только даты ────────────────
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            if hf:
                for para in hf.paragraphs:
                    t = para.text.strip()
                    if not t:
                        continue
                    if RE_DATE.search(t):
                        _set_para(para, RE_DATE.sub(date, t))

    return doc

# ─────────────────────────────────────────────────────────────
# ПРОЦЕССОР СЧЁТОВ
# ─────────────────────────────────────────────────────────────

def process_invoice(doc: Document, p: dict, amount: float) -> Document:
    amt_str   = format_amount(amount)
    amt_words = amount_to_words(amount)
    date      = p["date"]

    for tbl in doc.tables:
        nds_sum_col = -1
        if tbl.rows:
            for ci, hc in enumerate(tbl.rows[0].cells):
                if re.search(r"[Сс]умма\s*НДС|НДС\s*[Сс]умма", hc.text):
                    nds_sum_col = ci; break

        for ri, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                ct = cell.text.strip()
                if not ct:
                    continue

                # ── ИСПРАВЛЕНИЕ #2: Сохраняем "Номер документа" / "Дата составления" ──
                if "номер документа" in ct.lower() or "дата составления" in ct.lower():
                    for para in cell.paragraphs:
                        pt = _para_text(para)
                        if RE_DATE.search(pt):
                            _set_para(para, RE_DATE.sub(date, pt))
                    continue

                if re.search(r"22/122", ct):
                    for para in cell.paragraphs:
                        nf = re.sub(r"22/122%?", "Без НДС", _para_text(para))
                        _set_para(para, nf)
                elif RE_AMOUNT.search(ct):
                    if j == nds_sum_col and ri > 0:
                        for para in cell.paragraphs:
                            if RE_AMOUNT.search(_para_text(para)):
                                _set_para(para, "0,00")
                    else:
                        for para in cell.paragraphs:
                            nf = RE_AMOUNT.sub(amt_str, _para_text(para))
                            if nf != _para_text(para):
                                _set_para(para, nf)

    for para in doc.paragraphs:
        full = _para_text(para)
        nf   = full

        # Сохраняем "Номер документа" / "Дата составления"
        if "номер документа" in full.lower() or "дата составления" in full.lower():
            if RE_DATE.search(full):
                nf = RE_DATE.sub(date, full)
            if nf != full:
                _set_para(para, nf)
            continue

        if RE_DATE.search(full) and "счет" in full.lower():
            nf = RE_DATE.sub(date, nf)
        if RE_WORDS.search(nf):
            nf = RE_WORDS.sub(amt_words, nf)
            nf = re.sub(r",?\s*в\s+т\.?\s*ч\.?\s*НДС[^.]*", ", без НДС", nf)
            nf = re.sub(r"\s{2,}", " ", nf)
        nf = _remove_nds(nf)
        if nf != full:
            _set_para(para, nf)
    return doc

# ─────────────────────────────────────────────────────────────
# STREAMLIT UI
# ─────────────────────────────────────────────────────────────

def main():
    st.set_page_config(page_title="FRESH — Генератор ПВ", page_icon="🚗", layout="wide")
    st.title("🚗 FRESH — Генератор документов с ПВ")
    st.markdown("Автоматическое заполнение ДКП, ПКО и счетов при кредитных сделках")

    if "uploaded_files" not in st.session_state:
        st.session_state.uploaded_files = {}

    def auto_fill():
        if "inv1" in st.session_state.uploaded_files:
            try:
                doc  = Document(BytesIO(st.session_state.uploaded_files["inv1"]))
                data = extract_invoice_data(doc)
                mapping = {
                    "buyer_name": "ФИО",      "dkp_number": "Номер ДКП",
                    "date":       "Дата",     "car_brand":  "Марка",
                    "car_vin":    "VIN",      "car_color":  "Цвет",
                    "car_reg":    "Гос.номер",
                }
                filled = [lbl for k, lbl in mapping.items()
                          if data.get(k) and not st.session_state.get(k)]
                for k in mapping:
                    if data.get(k):
                        st.session_state[k] = data[k]
                if filled:
                    st.success(f"✅ Импортировано: {', '.join(filled)}")
                else:
                    st.warning("Данные не найдены — заполните вручную")
            except Exception as e:
                st.error(f"Ошибка: {e}")

    with st.sidebar:
        st.header("📁 Шаблоны документов")
        dkp_f  = st.file_uploader("ДКП",                           type=["docx"])
        pko_f  = st.file_uploader("ПКО",                           type=["docx"])
        inv1_f = st.file_uploader("Счёт №1  ⬅️ авто-заполнение",  type=["docx"])
        inv2_f = st.file_uploader("Счёт №2",                       type=["docx"])

        if dkp_f:  st.session_state.uploaded_files["dkp"]  = dkp_f.getvalue()
        if pko_f:  st.session_state.uploaded_files["pko"]  = pko_f.getvalue()
        if inv1_f: st.session_state.uploaded_files["inv1"] = inv1_f.getvalue()
        if inv2_f: st.session_state.uploaded_files["inv2"] = inv2_f.getvalue()

        st.markdown("---")
        st.markdown("**Статус загрузки:**")
        for k, lbl in [("dkp","ДКП"),("pko","ПКО"),("inv1","Счёт №1"),("inv2","Счёт №2")]:
            if k in st.session_state.uploaded_files:
                st.success(f"✅ {lbl}")
            else:
                st.warning(f"⚠️ {lbl} не загружен")

    col1, col2 = st.columns(2)

    defaults = {
        "buyer_name": "", "dkp_number": "",
        "date": datetime.today().strftime("%d.%m.%Y"),
        "car_brand": "", "car_vin": "",
        "car_color": "", "car_reg": "",
        "real_price": "", "pv_amount": "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    with col1:
        st.header("📋 Данные сделки")
        st.text_input("ФИО покупателя *",    key="buyer_name")
        st.text_input("Номер ДКП",            key="dkp_number")
        st.text_input("Дата документов *",    key="date", help="ДД.ММ.ГГГГ")
        st.text_input("Марка / Модель",       key="car_brand")
        st.text_input("VIN",                  key="car_vin")
        st.text_input("Цвет",                 key="car_color")
        st.text_input("Гос. номер",           key="car_reg")
        if "inv1" in st.session_state.uploaded_files:
            st.button("🔄 Авто-заполнить из Счёта №1",
                      on_click=auto_fill, use_container_width=True)

    with col2:
        st.header("💰 Суммы")
        st.text_input("Реальная цена авто (₽) *",      key="real_price", placeholder="3 250 000")
        st.text_input("Сумма искусственного ПВ (₽) *", key="pv_amount",  placeholder="250 000")

        real = parse_amount(st.session_state.real_price)
        pv   = parse_amount(st.session_state.pv_amount)
        if real > 0 and pv > 0:
            total = real + pv
            st.success(f"**Цена по документам (ДКП + Счёт №1)**\n\n{format_amount(total)} ₽")
            st.info(f"**Счёт №2 (к доплате покупателем)**\n\n{format_amount(real)} ₽")
            with st.expander("Суммы прописью"):
                st.write(f"Цена по документам: *{amount_to_words(total)}*")
                st.write(f"ПВ: *{amount_to_words(pv)}*")

    st.markdown("---")

    missing = []
    for k, lbl in [("dkp","ДКП"),("pko","ПКО"),("inv1","Счёт №1"),("inv2","Счёт №2")]:
        if k not in st.session_state.uploaded_files:
            missing.append(f"файл {lbl}")
    if not st.session_state.buyer_name.strip(): missing.append("ФИО покупателя")
    if parse_amount(st.session_state.real_price) <= 0: missing.append("реальная цена")
    if parse_amount(st.session_state.pv_amount)  <= 0: missing.append("сумма ПВ")

    if missing:
        st.warning("⚠️ Не заполнено: " + ", ".join(missing))

    if st.button("✅  Сгенерировать документы", type="primary",
                 use_container_width=True, disabled=bool(missing)):
        real  = parse_amount(st.session_state.real_price)
        pv    = parse_amount(st.session_state.pv_amount)
        total = real + pv
        params = {
            "buyer_name": st.session_state.buyer_name.strip(),
            "dkp_number": st.session_state.dkp_number.strip(),
            "date":       st.session_state.date.strip(),
            "car_brand":  st.session_state.car_brand.strip(),
            "car_vin":    st.session_state.car_vin.strip(),
            "car_color":  st.session_state.car_color.strip(),
            "car_reg":    st.session_state.car_reg.strip(),
            "real_price": real,
            "pv_amount":  pv,
            "new_price":  total,
        }
        surname   = (params["buyer_name"].split() or ["Клиент"])[0]
        date_safe = params["date"].replace(".", "-")

        tasks = [
            ("dkp",  f"ДКП_{surname}_{date_safe}.docx",
             lambda d, p: process_dkp(d, p)),
            ("pko",  f"ПКО_{surname}_{date_safe}.docx",
             lambda d, p: process_pko(d, p)),
            ("inv1", f"Счёт1_{surname}_{date_safe}.docx",
             lambda d, p: process_invoice(d, p, p["new_price"])),
            ("inv2", f"Счёт2_{surname}_{date_safe}.docx",
             lambda d, p: process_invoice(d, p, p["real_price"])),
        ]

        bar   = st.progress(0, text="Обрабатываю...")
        files = []
        for i, (key, fname, proc) in enumerate(tasks):
            if key in st.session_state.uploaded_files:
                try:
                    bar.progress(i / len(tasks), text=f"Обрабатываю {fname}…")
                    doc = Document(BytesIO(st.session_state.uploaded_files[key]))
                    doc = proc(doc, params)
                    buf = BytesIO(); doc.save(buf)
                    files.append((fname, buf.getvalue()))
                except Exception as e:
                    st.error(f"❌ {fname}: {e}")
            bar.progress((i + 1) / len(tasks))
        bar.empty()

        if files:
            st.success(f"✅ Готово — {len(files)} документа")
            cols = st.columns(len(files))
            for col, (fname, data) in zip(cols, files):
                with col:
                    st.download_button(
                        label=f"📄 {fname}",
                        data=data, file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)

if __name__ == "__main__":
    main()
